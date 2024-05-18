from docx import Document
from docx.shared import Inches
document = Document()

document.add_heading('簡単なWordドキュメントのタイトル', 0)
document.add_paragraph('簡単なWordドキュメントのテキスト')

def insert_image_into_docx(filename, image_path, width=Inches(1.5), height=None):
    try:
        doc = Document(filename)
        # 画像を挿入する場所を指定します（ここではドキュメントの末尾）
        doc.add_picture(image_path, width=width, height=height)
        # 変更を保存します
        doc.save("sample_answer.docx")
        print("画像が挿入されました:", image_path)
    except Exception as e:
        print("エラーが発生しました:", e)

def count_words_in_docx(filename):
    try:
        doc = Document(filename)
        total_words = 0
        for paragraph in doc.paragraphs:
            total_words += len(paragraph.text.split())
        return total_words
    except Exception as e:
        print("エラーが発生しました:", e)

# Wordファイル名を指定して文字数をカウント
filename = "sample.docx"  # 実際のファイル名に置き換えてください
image_filename = "example_image.jpg"
insert_image_into_docx(filename, image_filename)
word_count = count_words_in_docx(filename)
if word_count is not None:
    print(f"{filename} の文字数: {word_count}")

# document.save('sample_answer.docx')